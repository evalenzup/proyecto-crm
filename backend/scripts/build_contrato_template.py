"""
build_contrato_template.py
--------------------------
Convierte el machote 'CONTRATO DE SERVICIOSS.docx' en una plantilla docxtpl
insertando placeholders {{ }} en lugar de los valores concretos del ejemplo
(SIRONA HOSPITAL), preservando todo el formato, logo y membrete del Word.

La tabla de "personal asignado" se convierte en un loop {%tr ... %}.

Uso (dentro del contenedor backend, que tiene python-docx/docxtpl):
    python scripts/build_contrato_template.py ORIGEN.docx DESTINO.docx
"""
import sys
from docx import Document

# Reemplazos a nivel de run (case-insensitive). Ordenados de más largo a más corto
# para evitar reemplazos parciales. Cada valor concreto del machote → placeholder.
REPLACEMENTS = [
    # ── Prestador (NORTON / empresa) ──
    ("RODOLFO MUÑOZ BARBA", "{{ prestador_representante }}"),
    ("AIDA GARCÍA ORTEGA", "{{ prestador_propietario }}"),
    ("AIDA GARCIA ORTEGA", "{{ prestador_propietario }}"),
    ("GAOA611225II9", "{{ prestador_rfc }}"),
    ("GAOA-611225-II9", "{{ prestador_rfc }}"),
    ("A05-36064-10-6", "{{ prestador_registro_patronal }}"),
    ("AR8442/2021", "{{ prestador_repse_aviso }}"),
    ("05-02A039", "{{ prestador_licencia }}"),
    ("CITIBANAMEX", "{{ prestador_banco }}"),
    ("002022701317925713", "{{ prestador_clabe }}"),
    ("7013 1792571", "{{ prestador_cuenta }}"),
    # ── Cliente ──
    ("SIRONA HOSPITAL", "{{ cliente_razon_social }}"),
    ("ALEJANDRINA DIAZ VERGARA", "{{ cliente_representante }}"),
    ("Alejandrina Diaz Vergara", "{{ cliente_representante }}"),
    ("alejandrina diaz vergara", "{{ cliente_representante }}"),
    ("SHO1901292Z8", "{{ cliente_rfc }}"),
    ("Contabilidad.sirona@outlook.com", "{{ cliente_email }}"),
    # ── Contrato (precios, folio, vigencia, fecha firma) ──
    ("3,146.00", "{{ precio_fumigacion }}"),
    ("4,215.20", "{{ precio_sanitizacion }}"),
    ("6,488.00", "{{ precio_combo }}"),
    ("Septiembre 2025 al mes de Agosto de 2026", "{{ vigencia_texto }}"),
    ("Septiembre 2025 al mes de Agosto 2026", "{{ vigencia_texto }}"),
    ("Primero de Septiembre del Año Dos Mil Veinticinco", "{{ fecha_contrato_texto }}"),
    # Folio del certificado (5262) — siempre se refiere al mismo
    ("Nº 5262", "Nº {{ certificado_folio }}"),
    ("Nº 5262", "Nº {{ certificado_folio }}"),
    ("5262", "{{ certificado_folio }}"),
]

# Datos del trabajador de ejemplo → placeholders del loop (en celdas de la tabla)
TABLE_CELLS = {
    "SERGIO MOCTEZUMA HERNANDEZ": "{{ p.nombre }}",
    "25876814341": "{{ p.nss }}",
    "MOHS681007HBCCRR08": "{{ p.curp }}",
    "$419.88": "{{ p.salario }}",
    "419.88": "{{ p.salario }}",
    "TÉCNICO EN FUMIGACIÓN": "{{ p.puesto }}",
}


def _replace_in_runs(paragraph, pairs):
    """Reemplaza en cada run del párrafo (case-insensitive, preserva formato)."""
    for run in paragraph.runs:
        if not run.text:
            continue
        new = run.text
        low = new.lower()
        for needle, repl in pairs:
            if needle.lower() in low:
                # reemplazo case-insensitive simple
                idx = low.find(needle.lower())
                while idx != -1:
                    new = new[:idx] + repl + new[idx + len(needle):]
                    low = new.lower()
                    idx = low.find(needle.lower())
        if new != run.text:
            run.text = new


def _set_cell_text(cell, text):
    """Pone `text` en la celda preservando el formato del primer run; limpia el resto."""
    for p in cell.paragraphs:
        if not p.runs:
            continue
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
        # solo el primer párrafo con runs
        return
    # si no había runs, escribe en el primer párrafo
    if cell.paragraphs:
        cell.paragraphs[0].add_run(text)


def _control_row(data_row, tag_text):
    """Crea una fila clon de data_row con el tag de control en la 1a celda y el resto vacío."""
    import copy
    new_tr = copy.deepcopy(data_row._tr)
    from docx.table import _Row
    row = _Row(new_tr, data_row._parent)
    for i, cell in enumerate(row.cells):
        _set_cell_text(cell, tag_text if i == 0 else "")
    return row


def _merge_replace(paragraph, pairs):
    """Reemplazo a nivel de párrafo (une runs) — para valores fragmentados.
    Solo colapsa el formato si efectivamente hay un reemplazo pendiente."""
    if not paragraph.runs:
        return
    joined = "".join(r.text for r in paragraph.runs)
    new = joined
    low = new.lower()
    changed = False
    for needle, repl in pairs:
        nl = needle.lower()
        if nl in low:
            idx = low.find(nl)
            while idx != -1:
                new = new[:idx] + repl + new[idx + len(needle):]
                low = new.lower()
                idx = low.find(nl)
            changed = True
    if changed and new != joined:
        paragraph.runs[0].text = new
        for r in paragraph.runs[1:]:
            r.text = ""


def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def main():
    src, dst = sys.argv[1], sys.argv[2]
    doc = Document(src)

    # 1) Reemplazos a nivel de run (preserva formato; cubre valores contiguos)
    for p in _iter_paragraphs(doc):
        _replace_in_runs(p, REPLACEMENTS)

    # 2) Segunda pasada a nivel de párrafo para valores fragmentados
    #    (precios, vigencia, fecha de firma) que quedan partidos en varios runs.
    for p in _iter_paragraphs(doc):
        _merge_replace(p, REPLACEMENTS)

    # 2) Tabla de personal: localizar la que tiene encabezado de trabajador
    target_table = None
    for t in doc.tables:
        header = " ".join(c.text for c in t.rows[0].cells).upper()
        if "NOMBRE DEL TRABAJADOR" in header and "NSS" in header:
            target_table = t
            break

    if target_table is not None:
        rows = target_table.rows
        # rows[0] = encabezado; rows[1], rows[2] = trabajadores de ejemplo
        if len(rows) >= 2:
            data_row = rows[1]
            # Poner placeholders del loop en las celdas (orden: nombre, nss, curp, salario, puesto)
            cell_phs = ["{{ p.nombre }}", "{{ p.nss }}", "{{ p.curp }}", "{{ p.salario }}", "{{ p.puesto }}"]
            for i, cell in enumerate(data_row.cells):
                if i < len(cell_phs):
                    _set_cell_text(cell, cell_phs[i])
            # Patrón docxtpl: fila de control 'for' ANTES y 'endfor' DESPUÉS de la fila de datos
            for_row = _control_row(data_row, "{%tr for p in personal %}")
            endfor_row = _control_row(data_row, "{%tr endfor %}")
            data_row._tr.addprevious(for_row._tr)
            data_row._tr.addnext(endfor_row._tr)
            # Eliminar filas de datos sobrantes (el 2o trabajador de ejemplo)
            for extra in list(target_table.rows):
                txt = " ".join(c.text for c in extra.cells)
                if "GARCIA PALMA" in txt or "25038636285" in txt:
                    extra._tr.getparent().remove(extra._tr)

    doc.save(dst)
    print(f"Plantilla generada: {dst}")


if __name__ == "__main__":
    main()
